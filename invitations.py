#! python
# coding=utf-8
# invitations.py - Generate a Word document with custom invitations

import docx
doc = docx.Document('invitation.docx')
# 去除空行
if len(doc.paragraphs[0].runs) == 0:
    p = doc.paragraphs[0]._element
    p.getparent().remove(p)

# 写入文本
with open ('guests.txt') as f:
    lines = f.readlines()
    for line in lines:
        doc.add_paragraph('It would be a pleasure to have the company of', 'invitation1') 
        doc.add_paragraph(line.strip(), 'invitation2')
        doc.add_paragraph('at 11010 Memany of Lane at the evening of', 'invitation1')
        doc.add_paragraph('April 1st', 'invitation3')
        p = doc.add_paragraph("at 7 o'clock", 'invitation1')
        p.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.save ('newinvitation.docx')
